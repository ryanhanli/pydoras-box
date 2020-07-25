import docx

doc = docx.Document('demo.docx')
print(len(doc.paragraphs))
print(doc.paragraphs[0].text)
print(doc.paragraphs[1].text)
print(len(doc.paragraphs[1].runs))
print(doc.paragraphs[1].runs[0].text)
print(doc.paragraphs[1].runs[1].text)
print(doc.paragraphs[1].runs[2].text)
print(doc.paragraphs[1].runs[3].text)
print(doc.paragraphs[1].runs[4].text)
#print(doc.paragraphs[6].text) Double newline apparently

def getText(filename):
	doc = docx.Document(filename)
	fullText = []
	for paragraph in doc.paragraphs:
		fullText.append(paragraph.text)
	return '\n'.join(fullText)
	# return '\n\n'.join(fullText)
	# fullText.append('  ' + paragraph.text)
	
# Styles
print(" ")
print(doc.paragraphs[0].text)
print(doc.paragraphs[0].style)
doc.paragraphs[0].style = 'Normal'
print(doc.paragraphs[1].text)
print((doc.paragraphs[1].runs[0].text, doc.paragraphs[1].runs[1].text, doc.paragraphs[1].runs[2].text, doc.paragraphs[1].runs[3].text))
doc.paragraphs[1].runs[0].style = 'Quote Char'
doc.paragraphs[1].runs[1].underline = True
doc.paragraphs[1].runs[3].underline = True
doc.save('restyled.docx')

# Writing Word Documents
doc = docx.Document()
doc.add_paragraph('Hello world!', 'Title') # Adding style to this here.
paraObj1 = doc.add_paragraph('This is a second paragraph.') # Returns the paragraph
paraObj2 = doc.add_paragraph('This is a yet another paragraph.')
paraObj1.add_run(' This text is being added to the second paragraph.')
doc.save('helloworld.docx')
